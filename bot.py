from openai import OpenAI
from telegram import Update, BotCommand
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, filters, ContextTypes
from openai import OpenAI
import os

# Load secrets from environment variables 
TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "8292906924:AAHjgMruitNDsI-x6OPkstV5Pta-8p-TnGI")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "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")
client = OpenAI(api_key=OPENAI_API_KEY)
async def set_commands(app):
    commands = [
        BotCommand("start", "Start bot"),
        BotCommand("help", "Help"),
        BotCommand("scheme", "Generate scheme"),
        BotCommand("lessonplan", "Generate lesson plan"),
        BotCommand("notes", "Generate notes"),
        BotCommand("quiz", "Generate quiz"),
        BotCommand("timetable", "Generate timetable"),
    ]
    await app.bot.set_my_commands(commands)
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Welcome to Scheme Bot!\n\n"
        "/scheme\n"
        "/lessonplan\n"
        "/notes\n"
        "/quiz\n"
        "/timetable"
    )


async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Example:\n"
        "/scheme Grade 10 Physics Term 2"
    )


async def scheme(update: Update, context: ContextTypes.DEFAULT_TYPE):
    from docx import Document

    args = " ".join(context.args)

    if not args:
        await update.message.reply_text(
            "Use:\n/scheme Grade 10 Physics Term 2"
        )
        return

    doc = Document()

    doc.add_heading("SCHEME OF WORK", 0)
    doc.add_paragraph(f"{args}")

    table = doc.add_table(rows=1, cols=6)

    hdr = table.rows[0].cells
    hdr[0].text = "Week"
    hdr[1].text = "Lesson"
    hdr[2].text = "Strand"
    hdr[3].text = "Sub-Strand"
    hdr[4].text = "Learning Outcomes"
    hdr[5].text = "Activities"

    sample_data = [
        ["1", "1-4", "Mechanics", "Measurement", "Understand measurement concepts", "Class discussion"],
        ["2", "5-8", "Mechanics", "Motion", "Explain motion", "Practical examples"],
        ["3", "9-12", "Thermal Physics", "Temperature", "Measure temperature", "Lab activity"],
    ]

    for row in sample_data:
        cells = table.add_row().cells
        for i, item in enumerate(row):
            cells[i].text = item

    filename = "scheme_of_work.docx"
    doc.save(filename)

    await update.message.reply_document(document=open(filename, "rb"))


async def lessonplan(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Generating lesson plan...")


async def notes(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Generating notes...")


async def quiz(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Generating quiz...")


async def timetable(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Generating timetable...")
client = OpenAI(api_key=OPENAI_API_KEY)

# Start command
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Send me a subject like:\nForm 2 Biology Term 1\nand I will generate a scheme of work."
    )

# Message handler
async def generate_scheme(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_input = update.message.text

    prompt = f"""
    Generate a detailed Scheme of Work for:
    {user_input}

    Include:
    - Weekly breakdown
    - Topics
    - Subtopics
    - Learning outcomes
    - Simple teacher-friendly format
    """

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    result = response.choices[0].message.content

    await update.message.reply_text(result[:4000])  
async def timetable(update: Update, context: ContextTypes.DEFAULT_TYPE):
    from docx import Document
    from datetime import datetime, timedelta

    if len(context.args) < 3:
        await update.message.reply_text(
            "Use:\n/timetable Grade10A 8:00 40"
        )
        return

    class_name = context.args[0]
    start_time = context.args[1]
    lesson_duration = int(context.args[2])

    subjects = [
        "Physics",
        "Math",
        "English",
        "Biology",
        "Chemistry",
        "Geography",
        "History",
        "CRE"
    ]

    current = datetime.strptime(start_time, "%H:%M")

    times = []

    for _ in range(6):
        end = current + timedelta(minutes=lesson_duration)
        times.append(
            f"{current.strftime('%H:%M')}-{end.strftime('%H:%M')}"
        )
        current = end

    doc = Document()
    doc.add_heading(f"{class_name} WEEKLY TIMETABLE", 0)

    table = doc.add_table(rows=1, cols=6)

    hdr = table.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Monday"
    hdr[2].text = "Tuesday"
    hdr[3].text = "Wednesday"
    hdr[4].text = "Thursday"
    hdr[5].text = "Friday"

    for i, time in enumerate(times):
        row = table.add_row().cells
        row[0].text = time

        for j in range(1, 6):
            row[j].text = subjects[(i + j) % len(subjects)]

    filename = f"{class_name}_timetable.docx"
    doc.save(filename)

    await update.message.reply_document(
        document=open(filename, "rb")
    )
async def ask_ai(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_message = update.message.text

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a smart education assistant. "
                        "Answer all academic questions clearly. "
                        "You can help with schemes of work, lesson plans, "
                        "physics, mathematics, timetables, and general knowledge."
                    )
                },
                {
                    "role": "user",
                    "content": user_message
                }
            ]
        )

        reply = response.choices[0].message.content

        await update.message.reply_text(reply)

    except Exception as e:
        await update.message.reply_text(f"Error: {str(e)}")# Telegram message limit


# Run bot
app = ApplicationBuilder().token("8292906924:AAHBzK4u8zMbkRLwO2kAfBlFL_bU9v0L0VU").post_init(set_commands).build()

app.add_handler(CommandHandler("start", start))
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, generate_scheme))
app.add_handler(CommandHandler("help", help_command))
app.add_handler(CommandHandler("scheme", scheme))
app.add_handler(CommandHandler("lessonplan", lessonplan))
app.add_handler(CommandHandler("notes", notes))
app.add_handler(CommandHandler("quiz", quiz))
app.add_handler(CommandHandler("timetable", timetable))
app.add_handler(
    MessageHandler(
        filters.TEXT & ~filters.COMMAND,
        ask_ai
    )
)
print("Bot is running...")
app.run_polling()